#!/usr/bin/env python3
"""
Generate a professional Microsoft Word progress report for Milestone 1.
Content: WHAT IS DONE ONLY — no plans. Detailed, well-formatted, technical.
Requires: pip install python-docx

Usage:
  python scripts/generate_progress_report_docx.py
  python scripts/generate_progress_report_docx.py --output path/to/report.docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)


def set_cell_text(cell, text, bold=False, font_size=10):
    """Set cell text; optionally bold; font size in pt."""
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(font_size)
    r.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    """Add a table with header row bold."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[c], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < ncols:
                set_cell_text(table.rows[ri + 1].cells[ci], str(cell_text))
    if col_widths and len(col_widths) >= ncols:
        try:
            for ci, w in enumerate(col_widths):
                if ci < ncols:
                    table.columns[ci].width = Inches(w)
        except Exception:
            pass
    doc.add_paragraph()


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet paragraph. If bold_prefix given, bold that part only."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_code_para(doc, text, indent=0.25):
    """Add a paragraph with monospace font (code/technical block)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ----- Title -----
    title = doc.add_heading("Milestone 1 — Design System Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Progress Report — Work Completed to Date")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("23 February 2026")
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    doc.add_paragraph()

    # ----- Executive Summary -----
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report describes what has been delivered on the Milestone 1 Design System Agent. The system is a "
        "multi-agent pipeline that generates production-quality React/JSX UI components from natural language prompts, "
        "using switchable design system libraries (Untitled UI, Metafore, or both). All content below refers to "
        "implemented and working functionality only."
    )
    doc.add_paragraph(
        "Delivered: a 4-agent LangGraph pipeline (Orchestrator, Discovery, Generator, QA) with six library-aware "
        "tools, a RAG layer over design-system data, multi-library support (24 + 31 + 55 components), a ChatGPT-style "
        "chatbot UI with code pinning and live preview, and an MCP server that exposes the component library over "
        "streamable HTTP. The backend uses Claude Sonnet for code generation and GPT-4o-mini for classification and "
        "discovery; the frontend runs on port 3851 with SSE streaming and a collapsible thinking bar."
    )
    doc.add_paragraph()

    # ----- 1. What Is Delivered (Overview) -----
    doc.add_heading("1. What Is Delivered — Overview", level=1)
    add_table(
        doc,
        ["Area", "Delivered"],
        [
            ("Purpose", "Generate React/JSX components from prompts using design system catalog and tokens"),
            ("Pipeline", "4-agent LangGraph: Orchestrator → Discovery → Generator → QA Reviewer"),
            ("Tools", "6 tools in agent/tools.py: list_components, get_component_spec, get_design_tokens, preview_component, verify_quality, check_accessibility"),
            ("RAG", "Per-library vector index (agent/rag.py); OpenAI text-embedding-3-small; disk cache in design_system/.rag_cache/"),
            ("Design systems", "Untitled UI (24 components, blue), Metafore (31, purple), Both (55 merged)"),
            ("UI", "Chatbot on port 3851: library switcher, code pinning, ThinkingBar, pipeline viz, inline preview"),
            ("MCP server", "FastMCP server at project root: list_components, get_component_spec, get_design_tokens, resources; streamable HTTP"),
        ],
        col_widths=[1.5, 4.0],
    )

    # ----- 2. Architecture and Pipeline -----
    doc.add_heading("2. Architecture and Pipeline", level=1)

    doc.add_heading("2.1 Four Agents", level=2)
    add_table(
        doc,
        ["Agent", "File", "Model / Type", "Role", "Typical time"],
        [
            ("Orchestrator", "agent/orchestrator.py", "GPT-4o-mini (skipped when pre-classified)", "Classify request; route to generate/discover/review/chat; manage QA retry loop (max 2)", "~0–1 s"),
            ("Discovery", "agent/discovery.py", "GPT-4o-mini", "Single LLM call with pre-loaded catalog; returns component suggestions and Tailwind patterns", "~1–2 s"),
            ("Generator", "agent/generator.py", "Claude Sonnet (Anthropic)", "Single LLM call; writes complete React/JSX with library-specific Tailwind patterns", "~10–15 s"),
            ("QA Reviewer", "agent/reviewer.py + inline", "Rule-based (no LLM)", "Calls verify_quality and check_accessibility; returns PASS/FAIL; triggers retry on FAIL", "~0.1 s"),
        ],
        col_widths=[1.0, 1.4, 1.2, 1.8, 0.8],
    )
    doc.add_paragraph(
        "Pre-classification: chatbot/server.py runs _fast_classify(intent) with GPT-4o-mini once and passes the "
        "workflow (generate | discover | review | chat) into the pipeline, so the Orchestrator classify_node skips "
        "its LLM call. Discovery uses one GPT-4o-mini call with pre-loaded catalog; QA uses only rule-based tools."
    )

    doc.add_heading("2.2 Pipeline Flows", level=2)
    add_table(
        doc,
        ["User intent", "Flow"],
        [
            ("generate", "Pre-classify → Pipeline: Classify(skip) → Discovery → Generation → QA → Respond"),
            ("discover", "Pre-classify → Pipeline: Classify(skip) → Discovery → Respond"),
            ("review", "Pre-classify → Pipeline: Classify(skip) → QA → Respond"),
            ("chat", "Pre-classify → Direct GPT-4o-mini + RAG (no pipeline)"),
        ],
        col_widths=[0.9, 4.6],
    )

    doc.add_heading("2.3 Orchestrator State (OrchestratorState)", level=2)
    doc.add_paragraph("The LangGraph state passed between nodes has the following shape:")
    add_code_para(doc, "messages          # LangChain message history\nworkflow           # \"generate\" | \"discover\" | \"review\" | \"chat\" (pre-classified or from classify_node)\nuser_request        # Exact user message (stored by classify node)\ndiscovery_output   # Components + Tailwind patterns from Discovery\ngenerated_code      # React/JSX from Generator\nqa_result          # \"PASS\" | \"FAIL\" from QA\nretry_count        # QA retry counter (max 2)\nlibrary            # \"untitledui\" | \"metafore\" | \"both\" (active design system)")

    doc.add_heading("2.4 SSE Event Types (Backend → Frontend)", level=2)
    doc.add_paragraph("agent/server.py sends these event types over the stream:")
    add_table(
        doc,
        ["type", "Meaning"],
        [
            ("status", "Status text: e.g. \"Analyzing your request...\", \"Searching component library...\", \"Generating React code...\", \"Reviewing code quality...\", \"Preparing response...\""),
            ("thinking", "Discovery or generation LLM tokens (shown in collapsible ThinkingBar; not in chat)"),
            ("chunk", "Final response content only (shown in chat)"),
            ("done", "Stream complete"),
            ("error", "Error message"),
        ],
        col_widths=[0.8, 4.7],
    )
    doc.add_paragraph()

    # ----- 3. Tools and RAG -----
    doc.add_heading("3. Tools and RAG", level=1)

    doc.add_heading("3.1 Six Tools (Implementation)", level=2)
    add_table(
        doc,
        ["Tool", "File", "Type", "Purpose"],
        [
            ("list_components", "agent/tools.py", "MCP-equivalent", "List components from active library catalog JSON (set_active_library in classify_node)"),
            ("get_component_spec", "agent/tools.py", "MCP-equivalent", "Full spec + Tailwind patterns for one component"),
            ("get_design_tokens", "agent/tools.py", "MCP-equivalent", "Colors, typography, spacing, shadows from active library tokens JSON"),
            ("preview_component", "agent/tools.py", "Custom", "Save code as preview.html (React + Babel + Tailwind sandbox)"),
            ("verify_quality", "agent/tools.py", "Custom", "Rule-based: PascalCase, Tailwind, no imports, design-system compliance; score from 100"),
            ("check_accessibility", "agent/tools.py", "Custom", "Rule-based: semantic HTML, aria-labels, focus states, contrast; score from 100"),
        ],
        col_widths=[1.2, 1.0, 1.0, 2.3],
    )
    doc.add_paragraph(
        "When library=\"both\", catalogs and tokens are merged for list_components, get_component_spec, and get_design_tokens."
    )

    doc.add_heading("3.2 RAG System (agent/rag.py)", level=2)
    doc.add_paragraph(
        "In-memory vector store with per-library indexes. Chunks are built from catalog components, token categories, "
        "and documentation. OpenAI text-embedding-3-small produces embeddings; numpy stores vectors. Embeddings are "
        "cached to disk in design_system/.rag_cache/ as embeddings_{library}_{fingerprint}.npy with fingerprint-based "
        "invalidation. Query returns top-k chunks by cosine similarity."
    )
    add_table(
        doc,
        ["Aspect", "Detail"],
        [
            ("Index size", "~42 chunks (Untitled UI), ~48 (Metafore), ~90 (both)"),
            ("Injection", "chatbot/server.py _handle_direct_chat (chat path); agent/orchestrator.py respond_node (chat path) with library param"),
            ("API", "build_index(library), query(text, k=3, library)"),
        ],
        col_widths=[1.2, 4.3],
    )
    doc.add_paragraph()

    # ----- 4. Backend and API -----
    doc.add_heading("4. Backend and API", level=1)

    doc.add_heading("4.1 Chatbot Server (chatbot/server.py)", level=2)
    doc.add_paragraph("Port: 3851. ThreadingMixIn for concurrent requests. Environment loaded from .env via python-dotenv.")
    add_table(
        doc,
        ["Method", "Endpoint", "Body / Query", "Response"],
        [
            ("GET", "/api/health", "—", "{ ok, has_api_key }"),
            ("GET", "/api/catalog", "?library=untitledui|metafore|both", "{ tokens, catalog } (library-aware)"),
            ("POST", "/api/chat/stream", "{ message, history, intent, library }", "SSE: status + thinking + chunk + done"),
            ("POST", "/api/chat", "{ message, history }", "{ content } (non-streaming, GPT-4o)"),
            ("POST", "/api/preview", "{ code }", "Full HTML for iframe (Tailwind config)"),
            ("POST", "/api/generate", "{ prompt }", "{ code } (Claude Sonnet)"),
            ("POST", "/api/generate-variants", "{ prompt, count, keywords }", "{ variants } (Claude Sonnet)"),
        ],
        col_widths=[0.6, 1.4, 1.6, 2.0],
    )
    doc.add_paragraph(
        "When USE_LANGGRAPH=true: handle_stream parses message, history, intent, library → _fast_classify(intent) → "
        "if chat then _handle_direct_chat (GPT-4o-mini + RAG); else _handle_langgraph_stream(..., workflow=..., library=...) → agent.server.run_agent_stream()."
    )

    doc.add_heading("4.2 Environment (.env)", level=2)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ("OPENAI_API_KEY", "Required. RAG embeddings, classification, discovery, chat (GPT-4o-mini)"),
            ("ANTHROPIC_API_KEY", "Required for code generation. Claude Sonnet in generator, generate_code, generate_variants. Fallback to GPT-4o if missing"),
            ("USE_LANGGRAPH", "true = 4-agent pipeline; false or unset = direct Claude/GPT streaming fallback"),
        ],
        col_widths=[1.6, 4.0],
    )
    doc.add_paragraph()

    # ----- 5. Frontend and UX -----
    doc.add_heading("5. Frontend and UX", level=1)

    doc.add_heading("5.1 Stack and Entry", level=2)
    doc.add_paragraph(
        "Chatbot: React 18 (UMD CDN), Babel standalone, no build step. index.html loads React, Babel, marked.js, "
        "highlight.js. Main app in chatbot/app.jsx (~1300 lines). Styles in chatbot/styles.css. Preview: iframe with "
        "React + Babel + Tailwind CDN; preview.html is generated by preview_component tool."
    )

    doc.add_heading("5.2 Key React Components", level=2)
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ("App", "State: conversations, streaming, agentStep, codeFiles, thinkingContent, selectedLibrary"),
            ("ThinkingBar", "Collapsible bar; during gen shows \"Thinking...\"; after shows text summary of thinking content"),
            ("AgentPipeline", "Real-time pipeline: Classify → Discovery → Generate → QA → Respond"),
            ("Sidebar", "Conversation list + library selector (Untitled UI / Metafore / Both)"),
            ("CodePanel", "Saved code files; pin icon per file; pinned section at top"),
            ("InputArea", "Textarea + pinned context bar (chips for pinned files)"),
            ("Message", "User/assistant messages; pinned file tags on user messages"),
            ("PreviewModal / InlinePreview / MultiPreview", "Fullscreen, inline, or multi-variant preview"),
            ("FeaturesPanel", "Quick Actions, Variants, Catalog tabs (library-aware catalog fetch)"),
        ],
        col_widths=[1.4, 4.2],
    )

    doc.add_heading("5.3 Library Switcher", level=2)
    doc.add_paragraph(
        "selectedLibrary state: 'untitledui' | 'metafore' | 'both'. Sidebar has segmented control. Value is sent in "
        "streamChat() POST body, /api/catalog?library= query, and all pipeline calls. FeaturesPanel catalog tab "
        "refetches when library changes."
    )

    doc.add_heading("5.4 Code Pin / Attach System", level=2)
    add_bullet(doc, "Pin icon on each code file (hover); pinned files in \"PINNED\" section at top of CodePanel.")
    add_bullet(doc, "InputArea shows pinned-context-bar with chips; sendMessage prepends: \"IMPORTANT: You MUST use this pinned code as the base component to modify...\"")
    add_bullet(doc, "Intent vs message: clean intent for classification; full message (with pinned code + CODE_RULE) sent to pipeline.")
    add_bullet(doc, "Pinned tags shown on user messages; quick actions use pinRef (\"pinned component\" vs \"last UI component\"). CODE_RULE hidden from user, sent to backend.")

    doc.add_heading("5.5 Theme and Code Extraction", level=2)
    doc.add_paragraph(
        "Dark theme; violet accent (--accent: #7c3aed). Backgrounds --bg, --bg-sidebar, --bg-chat; text --text, --text-muted; borders --border. "
        "extractBestRunnableCodeBlock() scores blocks (root.render +50000, PascalCase +30000). extractAllRunnableCodeBlocks() for variants. "
        "History truncation: code blocks replaced with [code block omitted], max 500 chars."
    )
    doc.add_paragraph()

    # ----- 6. Design System Data -----
    doc.add_heading("6. Design System Data", level=1)

    doc.add_heading("6.1 Library File Mapping", level=2)
    add_table(
        doc,
        ["Library", "Catalog file", "Tokens file", "Components", "Brand"],
        [
            ("untitledui", "design_system/catalog.json", "design_system/tokens.json", "24", "Blue #1570EF"),
            ("metafore", "design_system/metafore_catalog.json", "design_system/metafore_tokens.json", "31", "Purple #7F56D9"),
            ("both", "Merged", "Merged", "55", "Both palettes"),
        ],
        col_widths=[1.0, 1.6, 1.6, 0.8, 1.6],
    )

    doc.add_heading("6.2 Catalog Structure", level=2)
    doc.add_paragraph(
        "Each component in catalog JSON: name, description, props, tailwind_pattern, variants. Untitled UI: Button, "
        "Input, Badge, Avatar, Card, Table, Tabs, Modal, Select, Checkbox, EmptyState, StatsCard, SearchInput, "
        "ProgressBar, Toggle, Radio, Textarea, Dropdown, Tooltip, Tag, Pagination, FileUpload, LoadingIndicator, "
        "Notification. Metafore: same structure; base (21), foundations (4), application (5), pages (1); e.g. Button, "
        "ButtonUtility, ButtonGroup, InputGroup, Label, MultiSelect, Combobox, Slider, Tags, Progress, FeaturedIcon, "
        "DotIcon, RatingBadge, RatingStars, HomeScreen, etc."
    )

    doc.add_heading("6.3 Tokens and coding_guidelines.md", level=2)
    doc.add_paragraph(
        "tokens.json / metafore_tokens.json: primary (blue vs purple), gray, success, error, warning; font Inter; "
        "shadows (5 vs 7 levels); radius; tailwindMapping (primary→blue vs primary→purple). coding_guidelines.md is "
        "injected into the Generator prompt: PascalCase, function components only, Tailwind only, accessibility "
        "(focus, semantic HTML, aria-labels), no imports (React/ReactDOM via CDN), component structure, variant rules."
    )
    doc.add_paragraph()

    # ----- 7. Generator and QA (Technical) -----
    doc.add_heading("7. Generator and QA — Technical Detail", level=1)

    doc.add_heading("7.1 Generator (agent/generator.py)", level=2)
    doc.add_paragraph(
        "Library-specific system prompts cached per library (_generation_prompt_cache). Tailwind patterns differ by library: "
        "e.g. primary button Untitled UI bg-blue-600 hover:bg-blue-700, Metafore bg-purple-600 hover:bg-purple-700; "
        "focus ring focus:ring-blue-500 vs focus:ring-purple-500; input focus focus:border-blue-500 vs focus:border-purple-500. "
        "When library=\"both\", both pattern sets are included in the prompt."
    )

    doc.add_heading("7.2 QA: verify_quality (Score from 100)", level=2)
    add_bullet(doc, "PascalCase component name (-15 if missing); className/Tailwind (-10); no inline styles (-5); root.render (-15); under 150 lines (-5); no imports (-10); no hardcoded hex (-5); no class components (-15).")
    add_bullet(doc, "Untitled UI compliance: Button/Input rounded-lg (-5); Cards rounded-xl (-5); non-standard colors (-5); font override (-5); table tbody divide-y divide-gray-200 (-5).")

    doc.add_heading("7.3 QA: check_accessibility (Score from 100)", level=2)
    add_bullet(doc, "Semantic HTML (-10); aria-label on icon buttons (-15); alt on images (-15); labels on inputs (-15); focus state classes (-10); color contrast (-5).")

    doc.add_paragraph("Verdict: PASS if score >= 70 and no errors. FAIL triggers retry (max 2 attempts).")
    doc.add_paragraph()

    # ----- 8. MCP Server -----
    doc.add_heading("8. MCP Server (Delivered)", level=1)
    doc.add_paragraph(
        "FastMCP server at project root (server.py). Exposes: tools list_components, get_component_spec, get_design_tokens, "
        "get_component_spec(component_name); resources design-system://tokens and design-system://components. "
        "Reads from design_system/catalog.json and design_system/tokens.json (Untitled UI). Transport: streamable HTTP "
        "(default port 8000). JSON response mode. Used so Cloud Code or other clients can query the component library "
        "and tokens over HTTP."
    )
    doc.add_paragraph()

    # ----- 9. How to Run and Directory Structure -----
    doc.add_heading("9. How to Run and Directory Structure", level=1)

    doc.add_heading("9.1 Run Commands", level=2)
    add_code_para(doc, "cd milestone1-agent\npip install -r requirements.txt\ncd chatbot\npython server.py   # → http://localhost:3851")
    doc.add_paragraph("On Windows, if port 3851 is in use: netstat -ano | Select-String \":3851\" then taskkill /F /PID <pid>.")

    doc.add_heading("9.2 Directory Structure (Key Paths)", level=2)
    add_code_para(doc,
        "milestone1-agent/\n"
        "  .env, .env.example, PROJECT_CONTEXT.md, coding_guidelines.md, requirements.txt\n"
        "  agent/          orchestrator.py, discovery.py, generator.py, reviewer.py, tools.py, rag.py, server.py\n"
        "  chatbot/        server.py, index.html, app.jsx, styles.css, preview.html\n"
        "  design_system/  catalog.json, tokens.json, metafore_catalog.json, metafore_tokens.json, .rag_cache/\n"
        "  server.py       # MCP server (project root)\n"
        "  docs/           architecture.md, pipeline-flow-diagram.png, *.excalidraw"
    )
    doc.add_paragraph()

    # ----- End -----
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("— End of report —")
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Calibri"

    return doc


def main():
    root = Path(__file__).resolve().parent.parent
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            docx_path = Path(sys.argv[idx + 1])
        else:
            docx_path = root / "docs" / "Milestone1_Progress_Report.docx"
    else:
        docx_path = root / "docs" / "Milestone1_Progress_Report.docx"

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    main()
