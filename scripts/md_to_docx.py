#!/usr/bin/env python3
"""
Convert MAKER_WIDGET_GENERATION_PLAN.md to a Word document (.docx).
Requires: pip install python-docx

Usage:
  python scripts/md_to_docx.py
  python scripts/md_to_docx.py --output path/to/custom.docx
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)


def parse_md_table(lines, start_i):
    """Parse a markdown table; return (list of rows, index after table)."""
    rows = []
    i = start_i
    while i < len(lines):
        line = lines[i]
        if not line.strip() or not "|" in line:
            break
        # Skip separator line (|---|---|)
        if re.match(r"^\s*\|[\s\-:]+\|", line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|") if c.strip() != ""]
        if not cells:
            break
        rows.append(cells)
        i += 1
    return rows, i


def strip_bold(s):
    """Replace **text** with plain text for table cells."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def add_paragraph_with_format(doc, text, bold_parts=None):
    """Add a paragraph; optionally bold specific parts (e.g. **word**)."""
    p = doc.add_paragraph()
    if not bold_parts:
        # Simple: replace **x** and add runs
        rest = text
        while True:
            m = re.search(r"\*\*(.+?)\*\*", rest)
            if not m:
                p.add_run(rest)
                break
            p.add_run(rest[: m.start()])
            p.add_run(m.group(1)).bold = True
            rest = rest[m.end() :]
    else:
        p.add_run(text)
    return p


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        # Code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # End of code block
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Table
        if line.strip().startswith("|") and "|" in line:
            rows, next_i = parse_md_table(lines, i)
            i = next_i
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(table.rows[ri].cells):
                        table.rows[ri].cells[ci].text = strip_bold(cell)
            # Bold header row
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            doc.add_paragraph()
            continue

        # Bullet list
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            rest = text
            while True:
                m = re.search(r"\*\*(.+?)\*\*", rest)
                if not m:
                    p.add_run(rest)
                    break
                p.add_run(rest[: m.start()])
                p.add_run(m.group(1)).bold = True
                rest = rest[m.end() :]
            i += 1
            continue

        # Nested bullet (  - )
        if re.match(r"^\s{2,}-\s", line):
            text = re.sub(r"^\s+-\s+", "", line)
            p = doc.add_paragraph(style="List Bullet 2")
            rest = text
            while True:
                m = re.search(r"\*\*(.+?)\*\*", rest)
                if not m:
                    p.add_run(rest)
                    break
                p.add_run(rest[: m.start()])
                p.add_run(m.group(1)).bold = True
                rest = rest[m.end() :]
            i += 1
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph (may contain **bold**)
        p = doc.add_paragraph()
        rest = line.strip()
        while True:
            m = re.search(r"\*\*(.+?)\*\*", rest)
            if not m:
                p.add_run(rest)
                break
            p.add_run(rest[: m.start()])
            p.add_run(m.group(1)).bold = True
            rest = rest[m.end() :]
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def main():
    root = Path(__file__).resolve().parent.parent
    md_path = root / "docs" / "MAKER_WIDGET_GENERATION_PLAN.md"
    if not md_path.exists():
        print(f"Not found: {md_path}")
        sys.exit(1)

    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            docx_path = Path(sys.argv[idx + 1])
        else:
            docx_path = root / "docs" / "MAKER_WIDGET_GENERATION_PLAN.docx"
    else:
        docx_path = root / "docs" / "MAKER_WIDGET_GENERATION_PLAN.docx"

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    md_to_docx(md_path, docx_path)


if __name__ == "__main__":
    main()
